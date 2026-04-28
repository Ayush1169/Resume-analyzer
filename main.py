from docx import Document
from groq import Groq
import cohere
import re

# ✅ FIX 1: Initialize Cohere client (add your API key here)
import os
from dotenv import load_dotenv

load_dotenv()

co = cohere.Client(os.getenv("COHERE_API_KEY"))

def read_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    return text


def generate_feedback(found, missing, role):
    prompt = f"""
    Candidate applied for {role}
    Skills found: {found}
    Missing skills: {missing}

    Give short professional feedback with:
    - Strength
    - Improvements
    - Verdict
    """

    try:
        # ✅ FIX 2: Removed duplicate except block
        response = co.chat(
            model="command-r-08-2024",
            message=prompt
        )
        return response.text.strip()

    except Exception as e:
        print("AI Error:", e)  # 👈 check your terminal for exact error
        return f"AI feedback not available. Error: {str(e)}"
    


def analyze_resume(file_path, role=None):
    text = ""

    try:
        if file_path.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""

        elif file_path.endswith(".docx"):
            text = read_docx(file_path)

        else:
            return [], 0, [], "Unsupported file format"

    except Exception as e:
        return [], 0, [], f"Error reading file: {e}"

    print("RAW TEXT:\n", text)

    # Normalize
    text_lower = text.lower()
    text_clean = re.sub(r'[^a-zA-Z0-9\s]', ' ', text_lower)
    text_no_space = text_clean.replace(" ", "")
    words = text_clean.split()

    skills = {
        "python": ["python"],
        "java": ["java"],
        "html": ["html", "html5"],
        "css": ["css", "css3"],
        "machine learning": ["machine learning", "machinelearning", "ml"],
        "data analysis": ["data analysis", "dataanalysis"],
        "react": ["react", "reactjs"],
        "node": ["node", "nodejs", "node js"],
        "express": ["express"],
        "mongodb": ["mongodb"],
        "javascript": ["javascript", "js"]
    }

    role_skills = {
        "fullstack": ["html", "css", "javascript", "react", "node", "express", "mongodb"],
        "frontend": ["html", "css", "javascript", "react"],
        "backend": ["node", "express", "mongodb", "javascript"],
        "ml engineer": ["python", "machine learning"],
        "data scientist": ["python", "machine learning", "data analysis"],
        "software engineer": ["python", "java", "javascript"]
    }

    found = []
    for skill, keywords in skills.items():
        for keyword in keywords:
            if (keyword in text_clean or
                keyword.replace(" ", "") in text_no_space or
                any(keyword in word for word in words)):
                found.append(skill)
                break

    found = list(set(found))

    if role and role.lower() in role_skills:
        target_skills = role_skills[role.lower()]
    else:
        target_skills = list(skills.keys())

    score = 0
    for skill in found:
        if skill in target_skills:
            score += 15
        else:
            score += 5

    missing = [skill for skill in target_skills if skill not in found]

    # ✅ FIX 3: Only call AI if role exists
    return found, score, missing, None